#!/usr/bin/env python3
"""Check a filled document against the template structure for completeness.
   Also cleans up extra/redundant punctuation in content.

Usage: 
  python3 docx_check.py <filled.docx> [--output output.docx]
  python3 docx_check.py <filled.docx> <template_structure.json> [--clean output.docx]
  python3 docx_check.py <filled.docx> --check-only
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx_utils import (
    is_placeholder,
    is_heading,
    clean_extra_punctuation,
    DocumentChecker,
    detect_template_type,
    TEMPLATE_TYPE_FEASIBILITY,
    TEMPLATE_TYPE_CHANGE
)


def get_section_content(doc, heading_text):
    """Extract content paragraphs under a specific heading."""
    paragraphs = list(doc.paragraphs)
    in_section = False
    content_lines = []
    for para in paragraphs:
        text = para.text.strip()
        style_name = (para.style.name or '').lower()
        if text == heading_text and ('heading' in style_name or 'title' in style_name):
            in_section = True
            continue
        if in_section:
            if 'heading' in style_name or 'title' in style_name:
                break
            if text:
                content_lines.append(text)
    return content_lines


def clean_document(docx_path, output_path):
    """Clean extra/redundant punctuation from content paragraphs."""
    from docx import Document
    
    doc = Document(docx_path)
    cleaned_count = 0

    for para in doc.paragraphs:
        if is_heading(para):
            continue

        text = para.text.strip()
        if not text:
            continue

        original = text
        text = clean_extra_punctuation(text)
        text = text.strip()

        if not text and original:
            for run in para.runs:
                run.text = ""
            cleaned_count += 1
        elif text != original:
            if para.runs:
                para.runs[0].text = text
                for run in para.runs[1:]:
                    run._r.getparent().remove(run._r)
            cleaned_count += 1

    doc.save(output_path)
    return {"status": "ok", "output": str(output_path), "cleaned": cleaned_count}


def check_document(docx_path, structure_path=None, check_only=False):
    """Check a filled document for completeness and issues."""
    from docx import Document
    
    doc = Document(docx_path)
    
    issues = []
    
    checker = DocumentChecker(doc)
    issues = checker.check_all()
    
    template_type = detect_template_type(docx_path)
    if template_type:
        issues.append({
            "type": "template_type",
            "severity": "info",
            "message": f"检测到文档类型: {template_type}"
        })
    
    if structure_path:
        structure = json.loads(Path(structure_path).read_text(encoding='utf-8'))
        flat_sections = structure.get("flat_sections", structure.get("sections", []))
        
        if isinstance(flat_sections, list) and flat_sections and "heading" in flat_sections[0]:
            sections = flat_sections
        else:
            def flatten(sections_list):
                result = []
                for s in sections_list:
                    result.append(s)
                    for sub in flatten(s.get("subsections", [])):
                        result.append(sub)
                return result
            sections = flatten(flat_sections)

        filled = 0
        empty = 0
        placeholder_only = 0

        for sec in sections:
            heading = sec["heading"]
            level = sec.get("level", 1)
            style = sec.get("style", "")

            if 'toc' in style.lower():
                continue

            content_lines = get_section_content(doc, heading)

            if not content_lines:
                empty += 1
                issues.append({
                    "type": "empty_section",
                    "severity": "error",
                    "heading": heading,
                    "level": level,
                    "message": f"章节未填写: {heading}"
                })
            elif all(is_placeholder(line) for line in content_lines):
                placeholder_only += 1
                issues.append({
                    "type": "placeholder_section",
                    "severity": "warning",
                    "heading": heading,
                    "level": level,
                    "message": f"章节仅含占位符: {heading}"
                })
            else:
                filled += 1

    error_count = sum(1 for i in issues if i.get("severity") == "error")
    warning_count = sum(1 for i in issues if i.get("severity") == "warning")
    
    return {
        "status": "ok" if error_count == 0 else "errors_found",
        "total_issues": len(issues),
        "errors": error_count,
        "warnings": warning_count,
        "issues": issues
    }


def main():
    parser = argparse.ArgumentParser(
        description="Check a filled Word document for completeness"
    )
    parser.add_argument("document", help="Input Word document (.docx)")
    parser.add_argument("structure", nargs="?", help="Template structure JSON (optional)")
    parser.add_argument("--output", "-o", help="Output file for cleaned document")
    parser.add_argument("--clean", action="store_true", help="Clean punctuation in output")
    parser.add_argument("--check-only", action="store_true", help="Only check, don't modify")
    
    args = parser.parse_args()
    
    doc_path = Path(args.document)
    if not doc_path.exists():
        print(f"ERROR: 文件不存在: {args.document}", file=sys.stderr)
        sys.exit(1)
    
    if args.clean or args.output:
        output_path = args.output or args.document.replace(".docx", "_cleaned.docx")
        result = clean_document(str(doc_path), output_path)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    if not args.check_only or args.structure:
        result = check_document(str(doc_path), args.structure, args.check_only)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    if args.check_only and not args.structure and not args.clean:
        result = check_document(str(doc_path))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        if result["errors"] > 0:
            sys.exit(1)


if __name__ == "__main__":
    main()
