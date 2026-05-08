#!/usr/bin/env python3
"""
Generate Type B (变更上线文档) by copying a .docx template and filling content.

NO format conversion. Direct .docx → .docx.
Preserves ALL template elements: images, tables, headers/footers, styles.

Usage:
  python3 docx_gen_change.py <template.docx> <content.json> <output.docx>
"""
import json
import sys
import os
import copy
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Maps content JSON key → exact template heading text
SECTION_MAP = {
    '部署架构图': '部署架构图（系统上线需要，一般变更不需要）',
    '变更计划': '变更计划简要说明:',
    '变更目的': '变更目的',
    '变更内容': '变更内容',
    '变更前准备': '变更前准备',
    '变更执行': '变更执行',
    '变更步骤': '变更步骤',
    '升级完成后检查': '升级完成后检查',
    '变更回退': '变更回退',
    '参数设置': '参数设置（可选）',
    '配置更新': '配置更新',
    '变更后工作': '变更后工作',
    '变更后检查': '变更后检查。',
}


def _find_heading_indices(doc):
    """Find paragraph indices for template section headings."""
    indices = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for key, heading in SECTION_MAP.items():
            if text == heading and heading not in indices:
                indices[heading] = i
                break
    return indices


def _section_has_content(doc, start_idx, end_idx):
    """Check if section has meaningful content (not just '。' or empty)."""
    for i in range(start_idx + 1, end_idx):
        if i >= len(doc.paragraphs):
            break
        text = doc.paragraphs[i].text.strip()
        if text and text != '。':
            return True
    return False


def _make_para(text, ref_para=None, bold=False, indent=True, gray_italic=False):
    """Create a paragraph XML element with Normal style."""
    p = OxmlElement('w:p')

    # Paragraph properties - use Normal style, not heading style
    pPr = OxmlElement('w:pPr')

    # Set style to Normal
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)

    # Copy spacing/indents from reference if available (but NOT the style)
    if ref_para is not None:
        ref_pPr = ref_para._element.find(qn('w:pPr'))
        if ref_pPr is not None:
            # Copy spacing
            ref_spacing = ref_pPr.find(qn('w:spacing'))
            if ref_spacing is not None:
                pPr.append(copy.deepcopy(ref_spacing))
            # Copy margins
            ref_ind = ref_pPr.find(qn('w:ind'))
            if ref_ind is not None:
                pPr.append(copy.deepcopy(ref_ind))

    # Justify alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'both')

    # First-line indent
    if indent:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), '200')
        ind.set(qn('w:firstLine'), '480')

    p.append(pPr)

    # Run with text - minimal formatting, inherit from Normal style
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Hint for east Asian text rendering
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rPr.append(rFonts)

    # Language
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    rPr.append(lang)

    if bold:
        rPr.append(OxmlElement('w:b'))
    if gray_italic:
        rPr.append(OxmlElement('w:i'))
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        rPr.append(color)

    r.insert(0, rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _make_empty_para(ref_para=None):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)
    return p


def _make_table(headers, rows):
    """Create a table XML element with header shading."""
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        borders.append(b)
    tblPr.append(borders)
    tbl.append(tblPr)

    grid = OxmlElement('w:tblGrid')
    for _ in headers:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(9000 // len(headers)))
        grid.append(col)
    tbl.append(grid)

    # Header row
    tr = OxmlElement('w:tr')
    for h in headers:
        tr.append(_make_cell(h, bold=True, shading='D9E2F3'))
    tbl.append(tr)

    # Data rows
    for row in rows:
        tr = OxmlElement('w:tr')
        for val in row:
            tr.append(_make_cell(str(val)))
        tbl.append(tr)
    return tbl


def _make_cell(text, bold=False, shading=None):
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    if shading:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), shading)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def _content_to_elements(content_list, ref_para):
    """Convert content lines to XML elements."""
    elems = []
    for line in content_list:
        if not line.strip():
            elems.append(_make_empty_para(ref_para))
            continue
        if line.startswith('【') and line.endswith('】'):
            elems.append(_make_para(line, ref_para, bold=True, indent=False))
            continue
        if line.startswith('步骤') and len(line) > 3 and line[3] in '0123456789':
            elems.append(_make_para(line, ref_para, bold=True))
            continue
        for label in ['操作：', '命令：', '预期：', '注意：', '提醒：', '检查：', '验证：']:
            if line.startswith(label):
                elems.append(_make_para(line, ref_para, bold=True))
                break
        else:
            if line.startswith('[截图:') or line.startswith('[截图：'):
                elems.append(_make_para(line, ref_para, gray_italic=True, indent=False))
            elif line.startswith('☐'):
                elems.append(_make_para(line, ref_para))
            else:
                elems.append(_make_para(line, ref_para))
    return elems


def _update_metadata_field(doc, label, value):
    """Update metadata value in table cell after label."""
    if not value:
        return
    # Search in tables first (template uses tables for metadata)
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if cell.text.strip().startswith(label.strip()) and i + 1 < len(row.cells):
                    target_cell = row.cells[i + 1]
                    if target_cell.paragraphs:
                        p = target_cell.paragraphs[0]
                        # Clear all runs, set one run with full value
                        for run in p.runs:
                            run.text = ''
                        if p.runs:
                            p.runs[0].text = value
                        else:
                            p.add_run(value)
                    return
    # Fallback: search in paragraphs
    skip = ['版', '编', '校', '部', '日', '类别', '编号']
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(label.strip()):
            if i + 1 < len(doc.paragraphs):
                nxt = doc.paragraphs[i + 1]
                txt = nxt.text.strip()
                if txt and not any(txt.startswith(s) for s in skip):
                    if nxt.runs:
                        nxt.runs[0].text = value
                    else:
                        nxt.add_run(value)
            break


def create_change_doc(template_path, content_json, output_path):
    """Copy template and fill content. NO conversion."""
    with open(content_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Copy template to output
    shutil.copy2(template_path, output_path)

    # Open the copy
    doc = Document(output_path)

    meta = data.get('metadata', {})
    project_name = data.get('project_name', '')

    # Update title (template has "某某系统变更方案", replace 某某 with project name)
    if project_name:
        for p in doc.paragraphs:
            if '系统变更方案' in p.text:
                for run in p.runs:
                    if '某某' in run.text:
                        run.text = run.text.replace('某某', project_name)
                break

    # Update metadata
    _update_metadata_field(doc, '版    本:', meta.get('version', ''))
    _update_metadata_field(doc, '编    制:', meta.get('author', ''))
    _update_metadata_field(doc, '校    对:', meta.get('reviewer', ''))
    _update_metadata_field(doc, '部    门:', meta.get('department', ''))

    # Update date in metadata table
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if '日    期' in cell.text:
                    if i + 1 < len(row.cells):
                        target = row.cells[i + 1]
                        if target.paragraphs and meta.get('date'):
                            p = target.paragraphs[0]
                            for run in p.runs:
                                run.text = ''
                            if p.runs:
                                p.runs[0].text = meta['date']
                            else:
                                p.add_run(meta['date'])

    # Find headings and fill content
    heading_indices = _find_heading_indices(doc)
    total = len(doc.paragraphs)
    ordered = sorted(heading_indices.items(), key=lambda x: x[1])

    insertions = []
    sections = data.get('sections', {})

    for content_key, template_heading in SECTION_MAP.items():
        section = sections.get(content_key, {})
        if not section:
            continue
        heading_idx = heading_indices.get(template_heading)
        if heading_idx is None:
            continue

        next_idx = total
        for h_text, h_idx in ordered:
            if h_idx > heading_idx:
                next_idx = h_idx
                break

        if _section_has_content(doc, heading_idx, next_idx):
            continue

        ref = doc.paragraphs[heading_idx]
        elems = []

        if 'steps_table' in section:
            sd = section['steps_table']
            headers = sd.get('headers', ['步骤', '操作内容', '预期结果', '截图'])
            rows = []
            for s in sd.get('steps', []):
                ops = s.get('ops', [s.get('op', '')])
                for k, op in enumerate(ops):
                    row = [s.get('step', '') if k == 0 else '', op]
                    if len(headers) > 2:
                        row.append(s.get('expected', '') if k == 0 else '')
                    if len(headers) > 3:
                        row.append(s.get('screenshot', '') if k == 0 else '')
                    rows.append(row[:len(headers)])
            elems.append(_make_table(headers, rows))
            elems.append(_make_empty_para(ref))
        elif 'table' in section:
            td = section['table']
            h = td.get('headers', [])
            r = td.get('rows', [])
            if h:
                elems.append(_make_table(h, r))
                elems.append(_make_empty_para(ref))
        elif 'content' in section:
            elems = _content_to_elements(section['content'], ref)

        if elems:
            insertions.append((heading_idx, elems))

    # Insert in reverse order (preserves indices)
    insertions.sort(key=lambda x: x[0], reverse=True)
    for heading_idx, elems in insertions:
        ref_element = doc.paragraphs[heading_idx]._element
        insert_after = ref_element
        for elem in elems:
            insert_after.addnext(elem)
            insert_after = elem

    doc.save(output_path)
    print(f'Change document saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 4:
        print(f'Usage: python3 {sys.argv[0]} <template.docx> <content.json> <output.docx>')
        sys.exit(1)
    create_change_doc(sys.argv[1], sys.argv[2], sys.argv[3])
