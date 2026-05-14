#!/usr/bin/env python3
"""Fill a Word template with content while preserving all formatting.

Usage: python3 docx_fill.py <template.docx> <content.json> <output.docx>

Content JSON format (from agent):
{
  "sections": {
    "Heading Text": {
      "content": ["Paragraph 1 text", "Paragraph 2 text"],
      "mode": "replace"  // or "append"
    }
  },
  "remove_placeholders": true
}

Key behaviors:
  - Preserves ALL existing formatting (fonts, alignment, spacing, tables, images)
  - New content paragraphs use the template's Normal style font settings
  - Strips outline numbering from content paragraphs (prevents heading-style numbering)
  - Heading numbering stays intact (only content under headings is modified)
"""

import argparse
import json
import re
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx_meta import update_meta
    from docx_utils import heading_match_score, make_table, make_empty_para, steps_table_to_rows
except ImportError:
    print("ERROR: python-docx not installed or docx_meta module not found. Run: pip3 install python-docx and ensure module is available.", file=sys.stderr)
    sys.exit(1)


PLACEHOLDER_PATTERNS = [
    re.compile(r'^\{.+\}$'),
    re.compile(r'^\[.+\]$'),
    re.compile(r'^【.+】$'),
    re.compile(r'^_+$'),
    re.compile(r'^…+$'),
    re.compile(r'^\.{3,}$'),
    re.compile(r'^-{3,}$'),
    re.compile(r'^={3,}$'),
    re.compile(r'^(请|待|TODO|FIXME|TBD|placeholder|填写|补充|内容)', re.IGNORECASE),
]


def is_placeholder(text):
    stripped = text.strip()
    if not stripped or len(stripped) < 2 or len(stripped) > 50:
        return False
    return any(p.match(stripped) for p in PLACEHOLDER_PATTERNS)


def prompt_string(prompt, default=""):
    if default:
        prompt_text = f"{prompt} [{default}]: "
    else:
        prompt_text = f"{prompt}: "
    try:
        value = input(prompt_text).strip()
        return value if value else default
    except EOFError:
        return default


def prompt_yes_no(prompt, default=True):
    default_hint = "Y/n" if default else "y/N"
    try:
        answer = input(f"{prompt} ({default_hint}): ").strip().lower()
    except EOFError:
        return default
    if not answer:
        return default
    return answer in ("y", "yes", "是", "yep")


def extract_template_fonts(doc):
    """Extract font settings from the template's Normal style and key heading styles.

    Returns a dict of style_name -> {rPr XML element} for font cloning.
    """
    fonts = {}
    for style in doc.styles:
        if style.type is not None and style.type.name == 'PARAGRAPH':
            name = style.name
            if name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
                rPr = style.element.find(qn('w:rPr'))
                if rPr is not None:
                    fonts[name] = deepcopy(rPr)
                # Also get paragraph format details
                pPr = style.element.find(qn('w:pPr'))
                if pPr is not None:
                    fonts[name + '_pPr'] = deepcopy(pPr)
    return fonts


def get_normal_rpr(template_fonts):
    """Get the Normal style's rPr for use in new content runs."""
    return template_fonts.get('Normal')


def strip_numbering(pPr_elem):
    """Remove or disable outline numbering (numPr) from paragraph properties.

    Sets numId=0 to explicitly suppress style-level numbering inheritance.
    """
    if pPr_elem is None:
        return
    existing = pPr_elem.find(qn('w:numPr'))
    if existing is not None:
        pPr_elem.remove(existing)
    # Explicitly set numId=0 to suppress any style-level numbering
    numPr = OxmlElement('w:numPr')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '0')
    numPr.append(numId)
    pPr_elem.append(numPr)


def clear_paragraph_content(para):
    """Remove all text from a paragraph while preserving its properties."""
    for run in para.runs:
        run.text = ""
    for elem in para._p.findall(qn('w:r')):
        for t in elem.findall(qn('w:t')):
            t.text = ""


def clone_paragraph_from(source_para, text="", force_normal=False, template_fonts=None):
    """Create a new paragraph XML element with proper formatting.

    Args:
        source_para: Source paragraph to clone style from
        text: Text content for the new paragraph
        force_normal: If True, override paragraph style to 'Normal'
        template_fonts: Dict from extract_template_fonts() for font settings

    Returns the new paragraph XML element.
    """
    new_p = OxmlElement('w:p')

    # Clone paragraph properties (alignment, spacing, indentation, etc.)
    pPr = source_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_pPr = deepcopy(pPr)
        # Strip outline numbering from content paragraphs
        strip_numbering(new_pPr)
        new_p.append(new_pPr)

    # Override style to Normal if requested
    if force_normal:
        new_pPr = new_p.find(qn('w:pPr'))
        if new_pPr is None:
            new_pPr = OxmlElement('w:pPr')
            new_p.insert(0, new_pPr)
        # Remove existing style
        pStyle = new_pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            new_pPr.remove(pStyle)
        # Set Normal style
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Normal')
        new_pPr.insert(0, pStyle)
        # Explicitly suppress outline numbering by setting numId=0
        existing_numPr = new_pPr.find(qn('w:numPr'))
        if existing_numPr is not None:
            new_pPr.remove(existing_numPr)
        numPr = OxmlElement('w:numPr')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '0')
        numPr.append(numId)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        new_pPr.append(numPr)

    # Build run with proper font settings
    r = OxmlElement('w:r')

    # Priority: template Normal style rPr > source paragraph's run rPr
    if template_fonts and 'Normal' in template_fonts:
        r.append(deepcopy(template_fonts['Normal']))
    elif source_para.runs:
        rPr = source_para.runs[0]._r.find(qn('w:rPr'))
        if rPr is not None:
            r.append(deepcopy(rPr))

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    return new_p


def apply_template_font_to_run(run, template_fonts):
    """Apply the template's Normal style font to an existing run.

    This ensures the run uses the correct font (especially east-asian fonts
        like 仿宋_GB2312 for Chinese text).
    """
    normal_rpr = get_normal_rpr(template_fonts)
    if normal_rpr is None:
        return

    existing_rPr = run._r.find(qn('w:rPr'))
    if existing_rPr is not None:
        # Merge: keep existing overrides, fill in missing from template
        template_rFonts = normal_rpr.find(qn('w:rFonts'))
        if template_rFonts is not None:
            existing_rFonts = existing_rPr.find(qn('w:rFonts'))
            if existing_rFonts is None:
                existing_rPr.insert(0, deepcopy(template_rFonts))
            else:
                # Fill in missing font attributes from template
                for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
                    qname = qn(f'w:{attr}')
                    if existing_rFonts.get(qname) is None:
                        val = template_rFonts.get(qname)
                        if val:
                            existing_rFonts.set(qname, val)
    else:
        run._r.insert(0, deepcopy(normal_rpr))


def find_section_range(doc, heading_text, fuzzy_threshold=0.8):
    """Find the paragraph index of a heading and the range of content after it.

    Supports fuzzy matching for headings that differ slightly in punctuation.

    Returns: (heading_idx, first_content_idx, end_idx)
    Where content paragraphs are [first_content_idx, end_idx)
    """
    paragraphs = list(doc.paragraphs)
    candidates = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        style_name = (para.style.name or '').lower()
        is_heading_style = 'heading' in style_name or 'title' in style_name
        score = heading_match_score(text, heading_text)
        if score >= fuzzy_threshold and is_heading_style:
            candidates.append((score, i))

    # Sort by score descending, pick best match
    if candidates:
        candidates.sort(key=lambda x: x[0], reverse=True)
        i = candidates[0][1]
        j = i + 1
        while j < len(paragraphs):
            s = (paragraphs[j].style.name or '').lower()
            if 'heading' in s or 'title' in s:
                break
            j += 1
        return i, i + 1, j

    return None, None, None


def fill_section(doc, heading_text, section_data, template_fonts):
    """Fill content into a section under the specified heading."""
    content_lines = section_data.get("content", [])
    mode = section_data.get("mode", "replace")

    heading_idx, content_start, content_end = find_section_range(doc, heading_text)

    if heading_idx is None:
        print(f"WARNING: Heading not found: '{heading_text}'", file=sys.stderr)
        return False

    heading_para = doc.paragraphs[heading_idx]
    content_indices = list(range(content_start, content_end))
    paragraphs = list(doc.paragraphs)

    if mode == "clear":
        for idx in content_indices:
            clear_paragraph_content(paragraphs[idx])
        _insert_section_tables(doc, heading_para, section_data)
        return True

    if mode == "append":
        insert_after_p = paragraphs[content_indices[-1]] if content_indices else heading_para
        for line in content_lines:
            new_p = clone_paragraph_from(heading_para, line, force_normal=True,
                                         template_fonts=template_fonts)
            insert_after_p._p.addnext(new_p)
            insert_after_p = type('P', (), {'_p': new_p})()
        _insert_section_tables(doc, heading_para, section_data)
        return True

    # "replace" mode
    fillable = []
    existing = []

    for idx in content_indices:
        text = paragraphs[idx].text.strip()
        if not text:
            fillable.append(idx)
        elif is_placeholder(text):
            fillable.append(idx)
        else:
            existing.append(idx)

    for idx in fillable:
        clear_paragraph_content(paragraphs[idx])

    template_para = paragraphs[content_indices[0]] if content_indices else heading_para

    insert_after_p = None
    filled_count = 0

    for i, line in enumerate(content_lines):
        if i < len(fillable):
            idx = fillable[i]
            p = paragraphs[idx]
            # Strip numbering from existing paragraphs too
            pPr = p._p.find(qn('w:pPr'))
            strip_numbering(pPr)

            if p.runs:
                p.runs[0].text = line
                for run in p.runs[1:]:
                    run._r.getparent().remove(run._r)
                # Ensure proper font
                apply_template_font_to_run(p.runs[0], template_fonts)
            else:
                run = p.add_run(line)
                apply_template_font_to_run(run, template_fonts)
            insert_after_p = p
            filled_count += 1
        else:
            if insert_after_p is None:
                insert_after_p = heading_para
            is_empty_section = (insert_after_p is heading_para)
            new_p = clone_paragraph_from(template_para, line,
                                          force_normal=is_empty_section,
                                          template_fonts=template_fonts)
            insert_after_p._p.addnext(new_p)
            insert_after_p = type('P', (), {'_p': new_p})()
            filled_count += 1

    # Handle table / steps_table insertion after text content
    _insert_section_tables(doc, heading_para, section_data)

    print(f"  Section '{heading_text}': filled {filled_count} paragraphs, "
          f"skipped {len(existing)} existing", file=sys.stderr)
    return True


def _insert_section_tables(doc, heading_para, section_data):
    """Insert table or steps_table into a section after text content."""
    insert_after = heading_para._element

    # Find the last content element under this heading to insert after
    heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p._element is heading_para._element:
            heading_idx = i
            break

    if heading_idx is not None:
        # Walk forward to find last non-heading paragraph before next heading
        paragraphs = list(doc.paragraphs)
        last_elem = heading_para._element
        for j in range(heading_idx + 1, len(paragraphs)):
            s = (paragraphs[j].style.name or '').lower()
            if 'heading' in s or 'title' in s:
                break
            # Skip empty paragraphs but track position
            last_elem = paragraphs[j]._element
        insert_after = last_elem

    if 'steps_table' in section_data:
        sd = section_data['steps_table']
        headers, rows = steps_table_to_rows(sd)
        if rows:
            tbl = make_table(headers, rows)
            insert_after.addnext(tbl)
            insert_after = tbl
            insert_after.addnext(make_empty_para())

    elif 'table' in section_data:
        td = section_data['table']
        headers = td.get('headers', [])
        rows = td.get('rows', [])
        if headers and rows:
            tbl = make_table(headers, rows)
            insert_after.addnext(tbl)
            insert_after = tbl
            insert_after.addnext(make_empty_para())


def fill_template(template_path, content_data, output_path):
    """Main entry point: fill template with content data."""
    doc = Document(template_path)
    template_fonts = extract_template_fonts(doc)
    sections = content_data.get("sections", {})
    filled_count = 0
    missing_count = 0

    print(f"Template fonts extracted: {list(template_fonts.keys())}", file=sys.stderr)

    for heading, section_data in sections.items():
        result = fill_section(doc, heading, section_data, template_fonts)
        if result:
            filled_count += 1
        else:
            missing_count += 1

    doc.save(output_path)
    print(f"\nDone: {filled_count} sections filled, {missing_count} headings not found", file=sys.stderr)
    print(f"Output saved to: {output_path}", file=sys.stderr)

    return {
        "filled": filled_count,
        "missing": missing_count,
        "output": str(output_path)
    }


def main():
    parser = argparse.ArgumentParser(description="Fill Word template with content")
    parser.add_argument("template", help="Path to the .docx template file")
    parser.add_argument("content", help="Path to the content JSON file")
    parser.add_argument("output", help="Path for the output .docx file")
    parser.add_argument("--author", help="Document author")
    parser.add_argument("--department", help="Document department or organization")
    parser.add_argument("--revision-date", help="Document date, e.g. 2026-05-12")
    parser.add_argument("--project-name", help="Project name used for homepage title smart rewrite")
    parser.add_argument("--title", help="Visible document title, overrides generated title")
    parser.add_argument("--no-interactive", action="store_true", help="Disable interactive metadata confirmation prompts")

    args = parser.parse_args()
    template_path = Path(args.template)
    content_path = Path(args.content)

    if not template_path.exists():
        print(f"ERROR: Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)
    if not content_path.exists():
        print(f"ERROR: Content file not found: {content_path}", file=sys.stderr)
        sys.exit(1)

    author = args.author or ""
    department = args.department or ""
    revision_date = args.revision_date or ""
    project_name = args.project_name or ""
    title = args.title or ""

    if sys.stdin.isatty() and not args.no_interactive:
        print("请先确认文档元信息，然后开始填充。")
        author = prompt_string("文档作者", author)
        department = prompt_string("部门/单位", department)
        revision_date = prompt_string("文档日期（YYYY-MM-DD）", revision_date)
        project_name = prompt_string("项目名称（用于首页标题智能修改）", project_name)
        title = prompt_string("文档标题（可选，留空则自动生成）", title)
        print("\n确认以下元信息：")
        print(f"  作者: {author}")
        print(f"  部门: {department}")
        print(f"  日期: {revision_date}")
        print(f"  项目名称: {project_name}")
        print(f"  文档标题: {title or '(自动生成)'}")
        if not prompt_yes_no("是否继续填充文档内容？", default=True):
            print("已取消。", file=sys.stderr)
            sys.exit(0)

    content_data = json.loads(content_path.read_text(encoding='utf-8'))
    result = fill_template(template_path, content_data, args.output)

    metadata_updates = {}
    if author:
        metadata_updates["author"] = author
    if department:
        metadata_updates["subject"] = department
    if revision_date:
        metadata_updates["revision_date"] = revision_date
    if project_name:
        metadata_updates["project_name"] = project_name
    if title:
        metadata_updates["title"] = title

    if metadata_updates:
        update_meta(args.output, args.output, metadata_updates)
        print(json.dumps({**result, "metadata_updates": metadata_updates}, ensure_ascii=False))
    else:
        print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
