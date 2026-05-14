#!/usr/bin/env python3
"""Shared utilities for docx-template-filler skill.

This module provides common functions used across all docx processing scripts.
"""

import re
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


PLACEHOLDER_PATTERNS = [
    re.compile(r'^\{.+\}$'),
    re.compile(r'^\[.+\]$'),
    re.compile(r'^【.+】$'),
    re.compile(r'^_+$'),
    re.compile(r'^…+$'),
    re.compile(r'^\.{3,}$'),
    re.compile(r'^-{3,}$'),
    re.compile(r'^={3,}$'),
    re.compile(r'^(请|待|TODO|FIXME|TBD|placeholder|填写|补充|内容)', re.IGNORECASE),
]

EXTRA_PUNCT_PATTERNS = [
    (re.compile(r'^[。，、；：！？…·\-_=]{1,3}$'), ''),
    (re.compile(r'^\s*[。，、]\s*$'), ''),
    (re.compile(r'。。+'), '。'),
    (re.compile(r'，，+'), '，'),
    (re.compile(r'。。'), '。'),
    (re.compile(r'\.\.\.(?!\.)'), '…'),
]


def is_placeholder(text: str) -> bool:
    """Check if text looks like a placeholder to be filled."""
    stripped = text.strip()
    if not stripped or len(stripped) < 2 or len(stripped) > 50:
        return False
    return any(p.match(stripped) for p in PLACEHOLDER_PATTERNS)


def is_meaningful_content(text: str) -> bool:
    """Check if text is meaningful content (not just whitespace/empty)."""
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) <= 2 and not any('\u4e00' <= c <= '\u9fff' for c in stripped):
        return False
    return True


def is_heading(para) -> bool:
    """Check if a paragraph is a heading (by style name)."""
    style_name = (para.style.name or '').lower()
    return 'heading' in style_name or 'title' in style_name or style_name.startswith('toc')


def get_heading_level(para) -> int:
    """Extract heading level number (1-9), default 1."""
    style_name = para.style.name or ''
    match = re.search(r'(\d)', style_name)
    return int(match.group(1)) if match else 1


def has_images(para) -> bool:
    """Check if a paragraph contains images."""
    for run in para.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            return True
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
            return True
    return False


def has_table_after(doc, para_index: int) -> bool:
    """Check if there's a table immediately after a paragraph."""
    body = doc.element.body
    children = list(body)
    if para_index + 1 < len(children):
        next_elem = children[para_index + 1]
        return next_elem.tag.endswith('}tbl')
    return False


def clean_extra_punctuation(text: str) -> str:
    """Clean extra/redundant punctuation from text."""
    for pattern, replacement in EXTRA_PUNCT_PATTERNS:
        text = pattern.sub(replacement, text)
    return text.strip()


def safe_load_docx(docx_path: str) -> Tuple[Document, str]:
    """Safely load a Word document with error handling.
    
    Returns:
        (Document, error_message) - if error, Document will be None
    """
    path = Path(docx_path)
    
    if not path.exists():
        return None, f"文件不存在: {docx_path}"
    
    if not path.suffix.lower() in ['.docx', '.doc']:
        return None, f"不支持的文件类型: {path.suffix}"
    
    try:
        doc = Document(docx_path)
        return doc, ""
    except Exception as e:
        return None, f"无法读取Word文档: {str(e)}"


def get_or_create_pPr(para) -> Optional[OxmlElement]:
    """Get or create paragraph properties element."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)
    return pPr


def create_paragraph_with_style(
    text: str,
    style: str = "Normal",
    bold: bool = False,
    italic: bool = False,
    indent_first_line: int = 0,
    alignment: str = "both"
) -> OxmlElement:
    """Create a paragraph XML element with specified style.
    
    Args:
        text: Text content
        style: Paragraph style name (default: Normal)
        bold: Bold text
        italic: Italic text
        indent_first_line: First line indent in characters (0 = none)
        alignment: Text alignment (left, center, right, both)
    
    Returns:
        OxmlElement representing the paragraph
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    
    if indent_first_line > 0:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLineChars'), str(indent_first_line * 100))
        ind.set(qn('w:firstLine'), str(indent_first_line * 240))
        pPr.append(ind)
    
    p.append(pPr)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rPr.append(rFonts)
    
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    rPr.append(lang)
    
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    
    return p


def extract_template_fonts(doc: Document) -> Dict:
    """Extract font settings from the template's Normal style and key heading styles.
    
    Returns a dict of style_name -> {rPr XML element} for font cloning.
    """
    fonts = {}
    for style in doc.styles:
        if style.type is not None and style.type.name == 'PARAGRAPH':
            name = style.name
            if name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
                rPr = style.element.find(qn('w:rPr'))
                if rPr is not None:
                    fonts[name] = deepcopy(rPr)
                pPr = style.element.find(qn('w:pPr'))
                if pPr is not None:
                    fonts[name + '_pPr'] = deepcopy(pPr)
    return fonts


def strip_numbering(pPr_elem) -> None:
    """Remove or disable outline numbering from paragraph properties."""
    if pPr_elem is None:
        return
    existing = pPr_elem.find(qn('w:numPr'))
    if existing is not None:
        pPr_elem.remove(existing)
    numPr = OxmlElement('w:numPr')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '0')
    numPr.append(numId)
    pPr_elem.append(numPr)


def find_screenshot_placeholders(doc: Document) -> List[Dict]:
    """Find all screenshot placeholder markers in the document.
    
    Returns:
        List of dicts with paragraph index and text
    """
    results = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '[截图' in text or '【截图' in text:
            results.append({
                "index": i,
                "text": text,
                "filled": not is_placeholder(text)
            })
    return results


def find_unfilled_placeholders(doc: Document) -> List[Dict]:
    """Find all unfilled placeholder paragraphs in the document.
    
    Returns:
        List of dicts with paragraph index and placeholder text
    """
    results = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if is_placeholder(text):
            results.append({
                "index": i,
                "text": text,
                "style": para.style.name
            })
    return results


def extract_header_footer_info(doc: Document) -> Dict:
    """Extract header and footer information from document.
    
    Returns:
        Dict with header/footer presence and content
    """
    result = {
        "has_header": False,
        "has_footer": False,
        "has_different_first_page": False,
        "header_content": [],
        "footer_content": [],
        "sections_info": []
    }
    
    # Check document section properties
    for section in doc.sections:
        section_info = {
            "start_type": str(section.start_type) if section.start_type else "continuous",
            "has_header": False,
            "has_footer": False,
            "different_first_page": section.different_first_page_header_footer
        }
        
        # Check header
        if section.header:
            section_info["has_header"] = True
            result["has_header"] = True
            for para in section.header.paragraphs:
                if para.text.strip():
                    result["header_content"].append(para.text.strip())
        
        # Check footer
        if section.footer:
            section_info["has_footer"] = True
            result["has_footer"] = True
            for para in section.footer.paragraphs:
                if para.text.strip():
                    result["footer_content"].append(para.text.strip())
        
        if section.different_first_page_header_footer:
            result["has_different_first_page"] = True
        
        result["sections_info"].append(section_info)
    
    return result


def copy_header_footer(src_doc: Document, dst_doc: Document) -> None:
    """Copy header and footer from source to destination document via XML deep-copy.

    Preserves all formatting, images, tabs, and field codes in headers/footers.

    Args:
        src_doc: Source document with headers/footers
        dst_doc: Destination document to copy to
    """
    if not src_doc.sections:
        return

    src_section = src_doc.sections[0]

    for dst_section in dst_doc.sections:
        # Copy header — deep copy XML paragraphs
        if src_section.header and not src_section.header.is_linked_to_previous:
            hdr_elem = dst_section.header._element
            # Remove existing paragraphs
            for child in list(hdr_elem):
                hdr_elem.remove(child)
            for para in src_section.header.paragraphs:
                hdr_elem.append(deepcopy(para._element))

        # Copy footer — deep copy XML paragraphs
        if src_section.footer and not src_section.footer.is_linked_to_previous:
            ftr_elem = dst_section.footer._element
            for child in list(ftr_elem):
                ftr_elem.remove(child)
            for para in src_section.footer.paragraphs:
                ftr_elem.append(deepcopy(para._element))

        # Copy section properties
        dst_section.different_first_page_header_footer = src_section.different_first_page_header_footer


def _copy_paragraph_format(src_para, dst_para) -> None:
    """Copy paragraph formatting from source to destination."""
    if src_para.paragraph_format.alignment:
        dst_para.paragraph_format.alignment = src_para.paragraph_format.alignment
    if src_para.paragraph_format.line_spacing:
        dst_para.paragraph_format.line_spacing = src_para.paragraph_format.line_spacing
    if src_para.paragraph_format.space_before:
        dst_para.paragraph_format.space_before = src_para.paragraph_format.space_before
    if src_para.paragraph_format.space_after:
        dst_para.paragraph_format.space_after = src_para.paragraph_format.space_after
    if src_para.paragraph_format.first_line_indent:
        dst_para.paragraph_format.first_line_indent = src_para.paragraph_format.first_line_indent


def extract_title_format(doc: Document) -> Dict:
    """Extract title format from first page.
    
    Returns:
        Dict with title format information
    """
    result = {
        "title_text": None,
        "title_font": {},
        "title_alignment": None,
        "title_level": None,
        "has_subtitle": False,
        "subtitle_text": None,
        "subtitle_font": {}
    }
    
    if not doc.paragraphs:
        return result
    
    first_para = doc.paragraphs[0]
    text = first_para.text.strip()
    
    if text:
        result["title_text"] = text
        
        if first_para.style.name:
            result["title_level"] = first_para.style.name
        
        if first_para.paragraph_format.alignment:
            result["title_alignment"] = str(first_para.paragraph_format.alignment)
        
        # Extract font info
        if first_para.runs:
            run = first_para.runs[0]
            font = run.font
            result["title_font"] = {
                "name": font.name.name if hasattr(font, 'name') and font.name else None,
                "size": str(font.size) if font.size else None,
                "bold": font.bold if font.bold is not None else False,
                "italic": font.italic if font.italic is not None else False,
                "underline": font.underline if font.underline else False
            }
    
    # Check for subtitle (second paragraph)
    if len(doc.paragraphs) > 1:
        second_para = doc.paragraphs[1]
        if second_para.text.strip():
            result["has_subtitle"] = True
            result["subtitle_text"] = second_para.text.strip()
            if second_para.runs:
                run = second_para.runs[0]
                font = run.font
                result["subtitle_font"] = {
                    "name": font.name.name if hasattr(font, 'name') and font.name else None,
                    "size": str(font.size) if font.size else None,
                    "bold": font.bold if font.bold is not None else False
                }
    
    return result


def match_title_format(target_doc: Document, source_format: Dict) -> List[Dict]:
    """Check if target document title matches source format.
    
    Args:
        target_doc: Document to check
        source_format: Title format from source (from extract_title_format)
        
    Returns:
        List of issues found
    """
    issues = []
    target_format = extract_title_format(target_doc)
    
    # Check title text exists
    if not target_format["title_text"]:
        issues.append({
            "type": "missing_title",
            "severity": "error",
            "message": "文档缺少标题"
        })
        return issues
    
    # Check title font
    if source_format.get("title_font"):
        src_font = source_format["title_font"]
        tgt_font = target_format.get("title_font", {})
        
        if src_font.get("name") and tgt_font.get("name"):
            if src_font["name"] != tgt_font["name"]:
                issues.append({
                    "type": "title_font_name",
                    "severity": "warning",
                    "message": f"标题字体不匹配: 期望 '{src_font['name']}', 实际 '{tgt_font['name']}'"
                })
        
        if src_font.get("size") and tgt_font.get("size"):
            if src_font["size"] != tgt_font["size"]:
                issues.append({
                    "type": "title_font_size",
                    "severity": "warning",
                    "message": f"标题字号不匹配: 期望 {src_font['size']}, 实际 {tgt_font['size']}"
                })
        
        if src_font.get("bold") != tgt_font.get("bold"):
            issues.append({
                "type": "title_font_bold",
                "severity": "warning",
                "message": f"标题粗体不匹配: 期望 {'粗体' if src_font['bold'] else '常规'}, 实际 {'粗体' if tgt_font.get('bold') else '常规'}"
            })
    
    # Check alignment
    if source_format.get("title_alignment") and target_format.get("title_alignment"):
        if source_format["title_alignment"] != target_format["title_alignment"]:
            issues.append({
                "type": "title_alignment",
                "severity": "warning",
                "message": "标题对齐方式不匹配"
            })
    
    return issues


def extract_paragraph_styles(doc: Document) -> Dict:
    """Extract all paragraph styles used in document.
    
    Returns:
        Dict mapping style names to format properties
    """
    styles = {}
    
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name and style_name not in styles:
            styles[style_name] = {
                "alignment": str(para.paragraph_format.alignment) if para.paragraph_format.alignment else None,
                "line_spacing": para.paragraph_format.line_spacing,
                "space_before": para.paragraph_format.space_before,
                "space_after": para.paragraph_format.space_after,
                "first_line_indent": str(para.paragraph_format.first_line_indent) if para.paragraph_format.first_line_indent else None,
                "keep_together": para.paragraph_format.keep_together,
                "keep_with_next": para.paragraph_format.keep_with_next
            }
    
    return styles


class DocumentChecker:
    """Comprehensive document checker for filled templates."""
    
    def __init__(self, doc: Document, structure: Optional[Dict] = None):
        self.doc = doc
        self.structure = structure or {}
        self.issues = []
    
    def check_all(self) -> List[Dict]:
        """Run all checks and return list of issues."""
        self.issues = []
        self._check_screenshot_placeholders()
        self._check_unfilled_placeholders()
        self._check_tables()
        self._check_images()
        return self.issues
    
    def _check_screenshot_placeholders(self):
        """Check for unfilled screenshot placeholders."""
        placeholders = find_screenshot_placeholders(self.doc)
        for ph in placeholders:
            if not ph['filled']:
                self.issues.append({
                    "type": "screenshot_placeholder",
                    "severity": "warning",
                    "index": ph['index'],
                    "text": ph['text'],
                    "message": f"截图占位符未填写: {ph['text']}"
                })
    
    def _check_unfilled_placeholders(self):
        """Check for unfilled text placeholders."""
        placeholders = find_unfilled_placeholders(self.doc)
        for ph in placeholders:
            self.issues.append({
                "type": "unfilled_placeholder",
                "severity": "warning",
                "index": ph['index'],
                "text": ph['text'],
                "style": ph['style'],
                "message": f"未填写占位符: {ph['text']}"
            })
    
    def _check_tables(self):
        """Check for empty or incomplete tables."""
        for i, table in enumerate(self.doc.tables):
            if not table.rows:
                self.issues.append({
                    "type": "empty_table",
                    "severity": "error",
                    "index": i,
                    "message": f"表格 {i+1} 为空"
                })
                continue
            if not table.rows[0].cells:
                self.issues.append({
                    "type": "empty_table",
                    "severity": "error", 
                    "index": i,
                    "message": f"表格 {i+1} 无表头"
                })
    
    def _check_images(self):
        """Check for images in document."""
        image_count = sum(1 for para in self.doc.paragraphs if has_images(para))
        if image_count == 0:
            self.issues.append({
                "type": "no_images",
                "severity": "info",
                "message": "文档中未发现图片"
            })
    
    def check_header_footer(self, reference_doc: Document = None, reference_info: Dict = None) -> List[Dict]:
        """Check header and footer presence and content.
        
        Args:
            reference_doc: Reference document to compare against
            reference_info: Pre-extracted reference header/footer info
            
        Returns:
            List of issues found
        }
        """
        issues = []
        
        # Extract current doc header/footer
        current_info = extract_header_footer_info(self.doc)
        
        # If reference provided, compare
        if reference_doc:
            ref_info = extract_header_footer_info(reference_doc)
        elif reference_info:
            ref_info = reference_info
        else:
            # Just report current state
            if not current_info["has_header"]:
                issues.append({
                    "type": "missing_header",
                    "severity": "warning",
                    "message": "文档缺少页眉"
                })
            if not current_info["has_footer"]:
                issues.append({
                    "type": "missing_footer",
                    "severity": "warning",
                    "message": "文档缺少页脚"
                })
            return issues
        
        # Compare header
        if ref_info["has_header"] and not current_info["has_header"]:
            issues.append({
                "type": "header_missing",
                "severity": "error",
                "message": "模板有页眉但生成的文档缺少页眉"
            })
        
        # Compare footer
        if ref_info["has_footer"] and not current_info["has_footer"]:
            issues.append({
                "type": "footer_missing",
                "severity": "error",
                "message": "模板有页脚但生成的文档缺少页脚"
            })
        
        # Check different first page setting
        if ref_info.get("has_different_first_page") != current_info.get("has_different_first_page"):
            issues.append({
                "type": "first_page_setting_mismatch",
                "severity": "warning",
                "message": "首页不同页眉页脚设置不匹配"
            })
        
        return issues
    
    def check_title_format(self, reference_doc: Document = None, reference_format: Dict = None) -> List[Dict]:
        """Check title format against reference.
        
        Args:
            reference_doc: Reference document
            reference_format: Pre-extracted title format
            
        Returns:
            List of issues found
        """
        issues = []
        
        # Get target format
        target_format = extract_title_format(self.doc)
        
        # Get source format
        if reference_doc:
            source_format = extract_title_format(reference_doc)
        elif reference_format:
            source_format = reference_format
        else:
            # Just check if title exists
            if not target_format["title_text"]:
                issues.append({
                    "type": "missing_title",
                    "severity": "error",
                    "message": "文档缺少标题"
                })
            return issues
        
        # Compare formats
        comparison_issues = match_title_format(self.doc, source_format)
        issues.extend(comparison_issues)
        
        return issues
    
    def check_paragraph_styles(self, reference_doc: Document = None) -> List[Dict]:
        """Check paragraph styles against reference.
        
        Args:
            reference_doc: Reference document to compare styles against
            
        Returns:
            List of issues found
        """
        issues = []
        
        current_styles = extract_paragraph_styles(self.doc)
        
        if reference_doc:
            ref_styles = extract_paragraph_styles(reference_doc)
            
            # Check for missing styles
            for style_name in ref_styles:
                if style_name not in current_styles:
                    issues.append({
                        "type": "missing_style",
                        "severity": "warning",
                        "message": f"缺少样式: {style_name}"
                    })
        
        return issues
    
    def check_format_comprehensive(self, reference_doc: Document = None) -> List[Dict]:
        """Run all format checks.
        
        Args:
            reference_doc: Reference document for comparison
            
        Returns:
            List of all issues found
        """
        self.issues = []
        
        # Basic content checks
        self._check_screenshot_placeholders()
        self._check_unfilled_placeholders()
        self._check_tables()
        self._check_images()
        
        # Format checks
        if reference_doc:
            self.issues.extend(self.check_header_footer(reference_doc))
            self.issues.extend(self.check_title_format(reference_doc))
            self.issues.extend(self.check_paragraph_styles(reference_doc))
        
        return self.issues


TEMPLATE_TYPE_FEASIBILITY = "feasibility"
TEMPLATE_TYPE_CHANGE = "change"

TEMPLATE_TYPES = {
    TEMPLATE_TYPE_FEASIBILITY: {
        "name": "可行性分析报告",
        "description": "项目立项可行性分析报告",
        "keywords": ["可行性", "分析", "报告", "立项", "项目"]
    },
    TEMPLATE_TYPE_CHANGE: {
        "name": "变更上线文档", 
        "description": "系统变更实施方案与操作指引",
        "keywords": ["变更", "上线", "方案", "实施", "回退"]
    }
}

TEMPLATE_SECTION_PATTERNS = {
    TEMPLATE_TYPE_FEASIBILITY: {
        "required": ["引言", "可行性分析前提", "项目背景", "需求分析", "技术方案", "项目计划", "资源配置", "效益分析", "结论与建议"],
        "optional": ["可行性分析前提", "修订记录", "项目实施", "风险分析", "项目管理", "测试计划"]
    },
    TEMPLATE_TYPE_CHANGE: {
        "required": ["变更目的", "变更内容", "变更步骤", "变更回退"],
        "optional": ["部署架构图", "变更计划", "变更前准备", "变更执行", "升级完成后检查", "参数设置", "配置更新", "软件及文档更新", "网络配置", "硬件配置", "变更后工作", "变更后检查"]
    }
}


def resolve_template(
    user_template_path: str = None,
    template_type: str = None,
    fallback_to_builtin: bool = True
) -> dict:
    """Resolve template path and type from user input or fallback to built-in.
    
    Priority:
    1. User provided template path -> analyze and use
    2. User provided template type -> use built-in template
    3. Auto-detect from content -> use built-in if matches
    4. Fallback to built-in templates
    
    Args:
        user_template_path: Path to user-provided template (optional)
        template_type: Requested template type - "feasibility" or "change" (optional)
        fallback_to_builtin: If True, use built-in templates when no user template
        
    Returns:
        dict with keys:
            - template_path: str - resolved template file path
            - template_type: str - detected/requested template type
            - is_user_template: bool - True if user provided the template
            - structure: dict - parsed template structure
            - source: str - "user", "builtin", or "auto_detected"
    """
    result = {
        "template_path": None,
        "template_type": None,
        "is_user_template": False,
        "structure": None,
        "source": None,
        "error": None
    }
    
    # Case 1: User provided a template path
    if user_template_path:
        path = Path(user_template_path)
        if not path.exists():
            result["error"] = f"用户提供的模板文件不存在: {user_template_path}"
            if fallback_to_builtin:
                return _fallback_to_builtin_template(template_type, result)
            return result
        
        # Detect template type from user template
        detected_type = detect_template_type(str(path))
        
        if template_type and template_type != detected_type:
            # User specified type but doesn't match - warn but proceed
            result["template_type"] = template_type
            result["warning"] = f"指定的类型 '{template_type}' 与检测到的类型 '{detected_type}' 不匹配"
        else:
            result["template_type"] = detected_type or template_type or TEMPLATE_TYPE_CHANGE
        
        result["template_path"] = str(path)
        result["is_user_template"] = True
        result["source"] = "user"
        
        # Parse structure if possible
        try:
            from docx import Document
            doc = Document(str(path))
            result["structure"] = _parse_template_structure(doc)
        except Exception as e:
            result["warning"] = f"无法解析模板结构: {e}"
        
        return result
    
    # Case 2: User specified template type (use built-in)
    if template_type:
        if template_type not in [TEMPLATE_TYPE_FEASIBILITY, TEMPLATE_TYPE_CHANGE]:
            result["error"] = f"未知的模板类型: {template_type}"
            if fallback_to_builtin:
                return _fallback_to_builtin_template(None, result)
            return result
        
        result["template_type"] = template_type
        result["template_path"] = get_template_path(template_type)
        result["source"] = "builtin_specified"
        
        if not Path(result["template_path"]).exists():
            result["error"] = f"内置模板不存在: {result['template_path']}"
            return result
        
        return result
    
    # Case 3: Auto-detect from context and use built-in
    if fallback_to_builtin:
        # Try feasibility first (more common)
        result["template_type"] = TEMPLATE_TYPE_CHANGE
        result["template_path"] = get_template_path(TEMPLATE_TYPE_CHANGE)
        result["source"] = "builtin_default"
        
        if Path(result["template_path"]).exists():
            return result
    
    result["error"] = "无法解析模板路径和类型"
    return result


def _fallback_to_builtin_template(preferred_type: str, result: dict) -> dict:
    """Fallback to built-in template."""
    template_type = preferred_type or TEMPLATE_TYPE_CHANGE
    result["template_path"] = get_template_path(template_type)
    result["template_type"] = template_type
    result["source"] = "builtin_fallback"
    result["is_user_template"] = False
    return result


def _parse_template_structure(doc) -> dict:
    """Parse template document structure.
    
    Returns:
        dict with sections, headings, and placeholders
    """
    sections = []
    headings = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = (para.style.name or '').lower()
        
        if is_heading(para):
            level = get_heading_level(para)
            headings.append({
                "text": text,
                "level": level,
                "style": para.style.name
            })
            sections.append({
                "heading": text,
                "level": level,
                "style": para.style.name,
                "placeholders": []
            })
        elif sections and is_placeholder(text):
            # Add placeholder to last section
            sections[-1]["placeholders"].append(text)
    
    return {
        "sections": sections,
        "headings": headings,
        "total_sections": len(sections),
        "total_headings": len(headings)
    }


def detect_template_type(docx_path: str) -> Optional[str]:
    """Detect the template type from document content.
    
    Returns:
        TEMPLATE_TYPE_FEASIBILITY, TEMPLATE_TYPE_CHANGE, or None if unclear
    """
    doc = Document(docx_path)
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    
    scores = {TEMPLATE_TYPE_FEASIBILITY: 0, TEMPLATE_TYPE_CHANGE: 0}
    
    for template_type, config in TEMPLATE_TYPES.items():
        for keyword in config["keywords"]:
            scores[template_type] += full_text.count(keyword)
    
    # Also check for section headings
    section_patterns = TEMPLATE_SECTION_PATTERNS
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_heading(para):
            for stype, patterns in section_patterns.items():
                for keyword in patterns.get("required", []) + patterns.get("optional", []):
                    if keyword in text:
                        scores[stype] += 2  # Heading match is stronger
    
    if scores[TEMPLATE_TYPE_FEASIBILITY] > scores[TEMPLATE_TYPE_CHANGE]:
        return TEMPLATE_TYPE_FEASIBILITY
    elif scores[TEMPLATE_TYPE_CHANGE] > scores[TEMPLATE_TYPE_FEASIBILITY]:
        return TEMPLATE_TYPE_CHANGE
    
    return None


def get_template_path(template_type: str) -> str:
    """Get the fixed template path for a given template type.
    
    Args:
        template_type: TEMPLATE_TYPE_FEASIBILITY or TEMPLATE_TYPE_CHANGE
        
    Returns:
        Absolute path to the template file
    """
    skill_dir = Path(__file__).parent.parent
    if template_type == TEMPLATE_TYPE_FEASIBILITY:
        return str(skill_dir / "templates" / "feasibility_template.docx")
    elif template_type == TEMPLATE_TYPE_CHANGE:
        return str(skill_dir / "templates" / "change_template.docx")
    else:
        raise ValueError(f"Unknown template type: {template_type}")


def prompt_template_type_choice() -> str:
    """Generate prompt for user to choose template type.
    
    Returns:
        Formatted prompt string for user interaction
    """
    return """请选择文档类型：
1. **可行性分析报告** — 项目立项可行性分析报告
2. **变更上线文档** — 系统变更实施方案与操作指引

请输入数字(1或2)或文档类型名称(可行性/变更)"""


def analyze_user_template(template_path: str) -> dict:
    """Analyze user-provided template and return structure info.
    
    Args:
        template_path: Path to user template
        
    Returns:
        dict with template analysis results
    """
    result = {
        "path": template_path,
        "exists": False,
        "type": None,
        "type_confidence": 0,
        "sections": [],
        "placeholders": [],
        "tables": [],
        "header_info": {},
        "footer_info": {},
        "has_different_first_page": False,
        "styles": {},
        "error": None
    }
    
    path = Path(template_path)
    if not path.exists():
        result["error"] = f"文件不存在: {template_path}"
        return result
    
    result["exists"] = True
    
    try:
        doc = Document(str(path))
        result["type"] = detect_template_type(str(path))
        result["sections"] = _parse_template_structure(doc)["sections"]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if is_placeholder(text):
                result["placeholders"].append(text)
        
        for ti, table in enumerate(doc.tables):
            table_info = {
                "index": ti,
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0,
                "has_data": len(table.rows) > 1
            }
            if table.rows and table.rows[0].cells:
                table_info["headers"] = [c.text.strip() for c in table.rows[0].cells]
            result["tables"].append(table_info)
        
        header_footer_info = extract_header_footer_info(doc)
        result["header_info"] = {
            "has_header": header_footer_info["has_header"],
            "content": header_footer_info["header_content"]
        }
        result["footer_info"] = {
            "has_footer": header_footer_info["has_footer"],
            "content": header_footer_info["footer_content"]
        }
        result["has_different_first_page"] = header_footer_info["has_different_first_page"]
        
        result["styles"] = extract_paragraph_styles(doc)
                
    except Exception as e:
        result["error"] = str(e)

    return result


def normalize_heading(text: str) -> str:
    """Normalize heading text for fuzzy matching.

    Strips whitespace, removes trailing punctuation (。.:：), and
    standardizes Chinese/English punctuation.
    """
    if not text:
        return ""
    text = text.strip()
    # Remove trailing punctuation commonly found in headings
    text = re.sub(r'[。\.\:：]+$', '', text)
    return text


def heading_match_score(heading_a: str, heading_b: str) -> float:
    """Compute a match score between two heading texts (0.0 - 1.0).

    Scoring:
        1.0  Exact match after normalization
        0.9  One contains the other after normalization
        0.8  One contains the other raw
        0.0  No match
    """
    if not heading_a or not heading_b:
        return 0.0
    a = normalize_heading(heading_a)
    b = normalize_heading(heading_b)
    if a == b:
        return 1.0
    if a in b or b in a:
        return 0.9
    # Fallback: raw containment (tolerates punctuation differences)
    raw_a = heading_a.strip()
    raw_b = heading_b.strip()
    if raw_a in raw_b or raw_b in raw_a:
        return 0.8
    return 0.0


def make_empty_para() -> OxmlElement:
    """Create an empty paragraph XML element with Normal style."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)
    return p


def make_cell(text: str, bold: bool = False, shading: Optional[str] = None) -> OxmlElement:
    """Create a table cell XML element."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    if shading:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), shading)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def make_table(headers: List[str], rows: List[List[str]]) -> OxmlElement:
    """Create a table XML element with header shading.

    Args:
        headers: List of header cell texts
        rows: List of row cell texts

    Returns:
        OxmlElement representing the table
    """
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        borders.append(b)
    tblPr.append(borders)
    tbl.append(tblPr)

    grid = OxmlElement('w:tblGrid')
    for _ in headers:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(9000 // len(headers)))
        grid.append(col)
    tbl.append(grid)

    # Header row
    tr = OxmlElement('w:tr')
    for h in headers:
        tr.append(make_cell(h, bold=True, shading='D9E2F3'))
    tbl.append(tr)

    # Data rows
    for row in rows:
        tr = OxmlElement('w:tr')
        for val in row:
            tr.append(make_cell(str(val)))
        tbl.append(tr)
    return tbl


def steps_table_to_rows(steps_table: Dict) -> Tuple[List[str], List[List[str]]]:
    """Convert a steps_table structure into headers and rows.

    Args:
        steps_table: {"headers": [...], "steps": [...]}

    Returns:
        (headers, rows)
    """
    headers = steps_table.get('headers', ['步骤', '操作内容', '预期结果', '截图'])
    rows = []
    for s in steps_table.get('steps', []):
        ops = s.get('ops', [s.get('op', '')])
        for k, op in enumerate(ops):
            row = [s.get('step', '') if k == 0 else '', op]
            if len(headers) > 2:
                row.append(s.get('expected', '') if k == 0 else '')
            if len(headers) > 3:
                row.append(s.get('screenshot', '') if k == 0 else '')
            rows.append(row[:len(headers)])
    return headers, rows
