#!/usr/bin/env python3
"""Read and update Word document metadata, headers, footers.
   Also handles first-line indent, TOC dirty marking, and version management.

Usage:
  python3 docx_meta.py read   <docx_file>
  python3 docx_meta.py update <docx_file> <output.docx> [options]
  python3 docx_meta.py indent <docx_file> <output.docx>
  python3 docx_meta.py toc    <docx_file> <output.docx>
  python3 docx_meta.py save   <docx_file> <workspace_dir> [--name base_name]
"""

import argparse
import glob
import json
import os
import re
import shutil
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def read_meta(docx_path):
    """Read all metadata from a Word document."""
    doc = Document(docx_path)
    props = doc.core_properties

    headers = []
    footers = []
    for i, section in enumerate(doc.sections):
        h = section.header
        f = section.footer
        h_texts = [p.text.strip() for p in h.paragraphs if p.text.strip()]
        f_texts = [p.text.strip() for p in f.paragraphs if p.text.strip()]
        # Check for images in header
        h_images = 0
        for p in h.paragraphs:
            h_images += len(p._p.findall('.//' + qn('w:drawing')))
            h_images += len(p._p.findall('.//' + qn('w:pict')))
        headers.append({
            "section": i,
            "linked_to_previous": h.is_linked_to_previous,
            "texts": h_texts,
            "has_images": h_images > 0,
            "image_count": h_images
        })
        footers.append({
            "section": i,
            "linked_to_previous": f.is_linked_to_previous,
            "texts": f_texts
        })

    result = {
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "comments": props.comments or "",
        "last_modified_by": props.last_modified_by or "",
        "revision": props.revision,
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "category": props.category or "",
        "keywords": props.keywords or "",
        "headers": headers,
        "footers": footers
    }
    return result


def update_meta(docx_path, output_path, updates):
    """Update document metadata, headers, and footers."""
    doc = Document(docx_path)
    props = doc.core_properties

    if "title" in updates and updates["title"]:
        props.title = updates["title"]
    if "author" in updates and updates["author"]:
        props.author = updates["author"]
    if "subject" in updates and updates["subject"]:
        props.subject = updates["subject"]
    if "comments" in updates and updates["comments"]:
        props.comments = updates["comments"]
    if "revision_date" in updates and updates["revision_date"]:
        try:
            dt = datetime.fromisoformat(updates["revision_date"])
            props.modified = dt
        except ValueError:
            for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日"]:
                try:
                    dt = datetime.strptime(updates["revision_date"], fmt)
                    props.modified = dt
                    break
                except ValueError:
                    continue

    # Update headers — preserve images, right-align text
    if "header_text" in updates and updates["header_text"]:
        _update_headers_with_images(doc, updates["header_text"])

    # Update footers (only if explicitly requested and no page number fields)
    if "footer_text" in updates and updates["footer_text"]:
        _update_footers(doc, updates["footer_text"])

    doc.save(output_path)
    return {"status": "ok", "output": str(output_path)}


def _extract_system_name(project_name):
    """Extract a short system name from the project name.

    Examples:
        '投资O32新增理财及信托产品投资功能' → '投资O32'
        '交易管理系统升级项目' → '交易管理'
        '某某系统建设' → '某某'

    Strategy: take characters up to the first Chinese verb/action keyword,
    or first 4-6 characters if no keyword found.
    """
    # Common action keywords that mark the boundary
    keywords = ['新增', '升级', '改造', '建设', '优化', '扩展', '开发', '实施', '部署']
    for kw in keywords:
        idx = project_name.find(kw)
        if idx > 0:
            return project_name[:idx]
    # Fallback: take first 4 characters
    return project_name[:min(4, len(project_name))]


def _read_template_header(docx_path):
    """Read the current header text from a document (to extract the template pattern)."""
    try:
        doc = Document(docx_path)
        for section in doc.sections:
            header = section.header
            if header.is_linked_to_previous:
                continue
            for para in header.paragraphs:
                text = para.text.strip()
                if text:
                    return text
    except Exception:
        pass
    return None


def _update_headers_with_images(doc, new_text):
    """Update header text while preserving the template's image+tab+text layout.

    Template header layout:
      [Image Run] [Tab Run] [Text Runs...]
    The tab stop positions text to the right. Image stays left.

    This function:
    1. Keeps image runs and the tab run intact
    2. Replaces only the text runs after the tab
    3. Preserves the tab stop definition (w:tabs) in pPr
    4. Preserves text run font formatting (宋体/微軟雅黑)
    """
    for section in doc.sections:
        header = section.header
        if header.is_linked_to_previous:
            has_content = any(p.text.strip() for p in header.paragraphs)
            if not has_content:
                continue

        for para in header.paragraphs:
            has_image = (
                para._p.findall('.//' + qn('w:drawing')) or
                para._p.findall('.//' + qn('w:pict'))
            )
            if not has_image:
                # No image — simple text-only header
                if para.text.strip() and para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run._r.getparent().remove(run._r)
                continue

            # Has image — use tab-based layout
            _update_header_with_tab_layout(para, new_text)


def _update_header_with_tab_layout(para, new_text):
    """Update a header paragraph that has image + tab + text layout.

    Layout: [image run] [tab run] [text runs...]
    Preserves image and tab, replaces text runs.
    """
    image_run = None
    tab_run = None
    text_runs = []
    text_rPr = None  # Font formatting from existing text runs

    for run in para.runs:
        r_elem = run._r
        has_drawing = (
            r_elem.findall('.//' + qn('w:drawing')) or
            r_elem.findall('.//' + qn('w:pict'))
        )
        has_tab = r_elem.findall(qn('w:tab'))

        if has_drawing:
            image_run = run
        elif has_tab:
            tab_run = run
        else:
            text_runs.append(run)
            if text_rPr is None:
                text_rPr = run._r.find(qn('w:rPr'))

    # Remove old text runs
    for run in text_runs:
        run._r.getparent().remove(run._r)

    # Build new text run with same font formatting
    new_run = OxmlElement('w:r')
    if text_rPr is not None:
        new_run.append(deepcopy(text_rPr))

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = new_text
    new_run.append(t)

    # Insert after the tab run (or image run if no tab)
    insert_after = tab_run or image_run
    if insert_after:
        insert_after._r.addnext(new_run)
    else:
        para._p.append(new_run)


def _update_footers(doc, new_text):
    """Update footer text, preserving page number fields."""
    for section in doc.sections:
        footer = section.footer
        if footer.is_linked_to_previous:
            has_content = any(p.text.strip() for p in footer.paragraphs)
            if not has_content:
                continue

        for para in footer.paragraphs:
            # Check for page number field codes
            instr_texts = para._p.findall('.//' + qn('w:instrText'))
            if instr_texts:
                continue  # Don't touch paragraphs with page number fields

            if para.text.strip():
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run._r.getparent().remove(run._r)
                else:
                    para.text = new_text


def apply_first_line_indent(docx_path, output_path, chars=2):
    """Apply first-line indent to all body paragraphs.

    Args:
        chars: Number of characters for first-line indent (default: 2 for Chinese)
    """
    doc = Document(docx_path)

    # firstLineChars: 100 = 1 character, 200 = 2 characters
    indent_val = chars * 100

    for para in doc.paragraphs:
        style_name = (para.style.name or '').lower()

        # Skip headings, title, TOC entries
        if 'heading' in style_name or 'title' in style_name or 'toc' in style_name:
            continue

        # Skip empty paragraphs
        if not para.text.strip():
            continue

        # Skip paragraphs that are part of tables
        if para._p.getparent() is not None:
            parent_tag = para._p.getparent().tag
            if parent_tag.endswith('}tc'):  # table cell
                continue

        # Apply first-line indent
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._p.insert(0, pPr)

        # Get or create indent element
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)

        # Remove existing firstLine/firstLineChars to avoid conflicts
        for attr in ['firstLine', 'firstLineChars']:
            if ind.get(qn(f'w:{attr}')) is not None:
                del ind.attrib[qn(f'w:{attr}')]

        # Set first-line indent (chars * 100)
        ind.set(qn('w:firstLineChars'), str(indent_val))

    doc.save(output_path)
    return {"status": "ok", "output": str(output_path), "indent": f"{chars} characters"}


def mark_toc_dirty(docx_path, output_path):
    """Mark all TOC fields as dirty so Word will prompt to update on open.

    Also updates PAGE/NUMPAGES field flags.
    """
    doc = Document(docx_path)

    # Find all field characters in the document
    body = doc.element.body
    field_chars = body.findall('.//' + qn('w:fldChar'))

    dirty_count = 0
    for fc in field_chars:
        fc_type = fc.get(qn('w:fldCharType'))
        if fc_type == 'begin':
            # Mark the field as dirty (needs update)
            fc.set(qn('w:dirty'), '1')
            dirty_count += 1

    # Also mark fields in headers and footers
    for section in doc.sections:
        for part in [section.header, section.footer]:
            for fc in part._element.findall('.//' + qn('w:fldChar')):
                if fc.get(qn('w:fldCharType')) == 'begin':
                    fc.set(qn('w:dirty'), '1')
                    dirty_count += 1

    doc.save(output_path)
    return {"status": "ok", "output": str(output_path), "fields_marked_dirty": dirty_count}


def version_save(docx_path, workspace_dir, base_name=None):
    """Save a versioned copy of the document with timestamp.

    Keeps only the last 3 versions. Saves to <workspace_dir>/doc/

    Returns:
        dict with save path and list of existing versions
    """
    doc_dir = os.path.join(workspace_dir, "doc")
    os.makedirs(doc_dir, exist_ok=True)

    # Generate filename
    if base_name is None:
        base_name = Path(docx_path).stem

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    versioned_name = f"{base_name}_v{timestamp}.docx"
    versioned_path = os.path.join(doc_dir, versioned_name)

    # Copy the file
    shutil.copy2(docx_path, versioned_path)

    # Find all versions of this document (use listdir for reliable Chinese filename matching)
    prefix = f"{base_name}_v"
    all_files = [f for f in os.listdir(doc_dir) if f.startswith(prefix) and f.endswith('.docx')]
    versions = sorted(
        [os.path.join(doc_dir, f) for f in all_files],
        key=os.path.getmtime,
        reverse=True
    )

    # Keep only last 3, delete older ones
    deleted = []
    if len(versions) > 3:
        for old_version in versions[3:]:
            os.remove(old_version)
            deleted.append(os.path.basename(old_version))
        versions = versions[:3]

    # Also save/update the "latest" symlink/copy
    latest_path = os.path.join(doc_dir, f"{base_name}_latest.docx")
    shutil.copy2(docx_path, latest_path)

    return {
        "status": "ok",
        "saved": versioned_path,
        "latest": latest_path,
        "versions": [os.path.basename(v) for v in versions],
        "deleted": deleted
    }


def main():
    parser = argparse.ArgumentParser(description="Word document metadata and formatting tools")
    subparsers = parser.add_subparsers(dest="action", required=True)

    # read
    p = subparsers.add_parser("read", help="Read document metadata")
    p.add_argument("document")

    # update
    p = subparsers.add_parser("update", help="Update document metadata/headers")
    p.add_argument("document")
    p.add_argument("output")
    p.add_argument("--title")
    p.add_argument("--author")
    p.add_argument("--subject")
    p.add_argument("--department")
    p.add_argument("--revision-date")
    p.add_argument("--header-text")
    p.add_argument("--project-name", help="Project name, auto-generates header as '<name>项目-可行性分析报告'")
    p.add_argument("--footer-text")
    p.add_argument("--comments")

    # indent
    p = subparsers.add_parser("indent", help="Apply first-line indent to body paragraphs")
    p.add_argument("document")
    p.add_argument("output")
    p.add_argument("--chars", type=int, default=2, help="Characters to indent (default: 2)")

    # toc
    p = subparsers.add_parser("toc", help="Mark TOC fields as dirty for update")
    p.add_argument("document")
    p.add_argument("output")

    # save (version management)
    p = subparsers.add_parser("save", help="Save versioned copy (keeps last 3)")
    p.add_argument("document")
    p.add_argument("workspace", help="Workspace directory (saves to <workspace>/doc/)")
    p.add_argument("--name", help="Base name for versioned file")

    args = parser.parse_args()

    if args.action == "read":
        print(json.dumps(read_meta(args.document), ensure_ascii=False, indent=2))

    elif args.action == "update":
        updates = {}
        if args.title: updates["title"] = args.title
        if args.author: updates["author"] = args.author
        if args.subject: updates["subject"] = args.subject
        if args.department and "subject" not in updates:
            updates["subject"] = args.department
        if args.revision_date: updates["revision_date"] = args.revision_date
        if args.header_text:
            updates["header_text"] = args.header_text
        elif args.project_name:
            # Auto-generate header from project name
            # Pattern: extract the suffix after '某某' in template header
            # e.g. template="某某功能项目-可行性分析报告"
            #       + project="投资O32新增理财及信托产品投资功能"
            #       → header="投资O32新增理财及信托产品投资功能项目-可行性分析报告"
            template_header = _read_template_header(args.document)
            if template_header:
                # Build header from template pattern.
                # Template: "某某系统某某功能项目-可行性分析报告"
                #   1st '某某' = system short name
                #   2nd '某某' = project name (before '功能')
                #
                # Strategy: replace only the SECOND '某某' with project name,
                # keep the 1st as the system name derived from project name.
                #
                # Example:
                #   project = "投资O32新增理财及信托产品投资功能"
                #   → header = "投资O32系统投资O32新增理财及信托产品投资项目-可行性分析报告"
                pname = args.project_name
                # Extract system short name: take first meaningful segment
                # (e.g. "投资O32" from "投资O32新增理财及信托产品投资功能")
                sys_name = _extract_system_name(pname)
                # Strip trailing '功能' from project name for the 2nd placeholder
                pname_short = pname.rstrip('功能') if pname.endswith('功能') else pname
                # Replace: 1st '某某' → system name, 2nd '某某功能' → project + '功能'
                # Use count=1 to replace one at a time
                header_text = template_header
                # Find positions of '某某'
                idx1 = header_text.find('某某')
                if idx1 >= 0:
                    idx2 = header_text.find('某某', idx1 + 2)
                    if idx2 >= 0:
                        # Replace 2nd first (longer match), then 1st
                        header_text = (
                            header_text[:idx2]
                            + pname_short + '功能'
                            + header_text[idx2 + 4:]  # skip '某某功能' (4 chars)
                        )
                        # Recalculate 1st position (unchanged since we replaced after it)
                        header_text = (
                            header_text[:idx1]
                            + sys_name
                            + header_text[idx1 + 2:]  # skip '某某' (2 chars)
                        )
                    else:
                        # Only one '某某' — replace with project name
                        header_text = header_text.replace('某某', pname, 1)
                updates["header_text"] = header_text
            else:
                updates["header_text"] = f"{args.project_name}项目-可行性分析报告"
        if args.footer_text: updates["footer_text"] = args.footer_text
        if args.comments: updates["comments"] = args.comments
        print(json.dumps(update_meta(args.document, args.output, updates), ensure_ascii=False))

    elif args.action == "indent":
        print(json.dumps(apply_first_line_indent(args.document, args.output, args.chars), ensure_ascii=False))

    elif args.action == "toc":
        print(json.dumps(mark_toc_dirty(args.document, args.output), ensure_ascii=False))

    elif args.action == "save":
        print(json.dumps(version_save(args.document, args.workspace, args.name), ensure_ascii=False))


if __name__ == "__main__":
    main()
