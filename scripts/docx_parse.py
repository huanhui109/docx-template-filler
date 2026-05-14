#!/usr/bin/env python3
"""Parse a Word template to extract its heading structure and content areas.

Usage: python3 docx_parse.py <template.docx> [--output structure.json]

Output JSON structure:
{
  "source": "template.docx",
  "title": "Document Title",
  "sections": [
    {
      "heading": "Heading Text",
      "level": 1,
      "style": "Heading 1",
      "content": [
        {"index": 0, "text": "Existing paragraph text", "style": "Normal", "is_placeholder": false}
      ],
      "tables": 0,
      "subsections": [...]
    }
  ],
  "metadata": {
    "total_headings": 5,
    "total_paragraphs": 42,
    "has_tables": true,
    "has_images": false
  }
}
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


# Common placeholder patterns
PLACEHOLDER_PATTERNS = [
    re.compile(r'^\{.+\}$'),           # {content}
    re.compile(r'^\[.+\]$'),           # [content]
    re.compile(r'^【.+】$'),            # 【内容】
    re.compile(r'^_+$'),               # ___
    re.compile(r'^…+$'),               # ……
    re.compile(r'^\.{3,}$'),           # ...
    re.compile(r'^-{3,}$'),            # ---
    re.compile(r'^={3,}$'),            # ===
    re.compile(r'^(请|待|TODO|FIXME|TBD|placeholder|填写|补充|内容)', re.IGNORECASE),
]


def _extract_system_name(project_name):
    """Extract a short system name from the project name.

    Examples:
        '投资O32新增理财及信托产品投资功能' → '投资O32'
        '交易管理系统升级项目' → '交易管理'
        '某某系统建设' → '某某'

    Strategy: take characters up to the first Chinese verb/action keyword,
    or first 4-6 characters if no keyword found.
    """
    # Common action keywords that mark the boundary
    keywords = ['新增', '升级', '改造', '建设', '优化', '扩展', '开发', '实施', '部署']
    for kw in keywords:
        idx = project_name.find(kw)
        if idx > 0:
            return project_name[:idx]
    # Fallback: take first 4 characters
    return project_name[:min(4, len(project_name))]


def _find_first_title_paragraph(doc):
    """Find the first title-like paragraph on the first page."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or '').lower()
        if 'title' in style_name or 'heading 1' in style_name or style_name.startswith('heading'):
            return para
    return None


def _read_template_header(doc):
    """Read the current header text from a document (to extract the template pattern)."""
    try:
        for section in doc.sections:
            header = section.header
            if header.is_linked_to_previous:
                continue
            for para in header.paragraphs:
                text = para.text.strip()
                if text:
                    return text
    except Exception:
        pass
    return None


def _replace_project_placeholder(template_text, project_name):
    """Replace placeholder text in a template title or header using the project name."""
    if '某某' not in template_text:
        return template_text

    count = template_text.count('某某')
    if count == 1:
        return template_text.replace('某某', project_name)

    sys_name = _extract_system_name(project_name)
    pname_short = project_name[len(sys_name):] if project_name.startswith(sys_name) else project_name
    if pname_short.endswith('功能'):
        pname_short = pname_short[:-1]

    result = template_text
    result = result.replace('某某', pname_short, 1)
    result = result.replace('某某', sys_name, 1)
    return result


def _generate_title_from_project_name(doc, project_name):
    """Generate a smart title for the document based on the project name and template."""
    if not project_name:
        return project_name

    title_para = _find_first_title_paragraph(doc)
    title_text = title_para.text.strip() if title_para is not None else ''
    header_text = _read_template_header(doc)

    if title_text and '某某' in title_text:
        return _replace_project_placeholder(title_text, project_name)

    if header_text and '某某' in header_text:
        return _replace_project_placeholder(header_text, project_name)

    if title_text:
        if project_name not in title_text:
            suffix_match = re.search(r'(项目.*|.*报告|.*方案|.*建议书)$', title_text)
            if suffix_match:
                return f"{project_name}{suffix_match.group(1)}"
            if title_text in ('可行性分析报告', '项目可行性分析报告', '变更方案'):
                return f"{project_name}{title_text}"
        return title_text

    return project_name


def is_heading(para):
    """Check if a paragraph is a heading (by style name)."""
    style_name = (para.style.name or '').lower()
    return 'heading' in style_name or style_name.startswith('toc')


def get_heading_level(para):
    """Extract heading level number (1-9), default 1."""
    style_name = para.style.name or ''
    match = re.search(r'(\d)', style_name)
    return int(match.group(1)) if match else 1


def is_placeholder(text):
    """Check if text looks like a placeholder to be filled."""
    stripped = text.strip()
    if not stripped or len(stripped) < 2:
        return False
    if len(stripped) > 50:
        return False
    return any(p.match(stripped) for p in PLACEHOLDER_PATTERNS)


def is_meaningful_content(text):
    """Check if text is meaningful content (not just whitespace/empty)."""
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) <= 2 and not any('\u4e00' <= c <= '\u9fff' for c in stripped):
        return False
    return True


def has_images(para):
    """Check if a paragraph contains images."""
    for run in para.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            return True
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
            return True
    return False


def has_table_after(doc, para_index):
    """Check if there's a table immediately after a paragraph."""
    body = doc.element.body
    children = list(body)
    if para_index + 1 < len(children):
        next_elem = children[para_index + 1]
        return next_elem.tag.endswith('}tbl')
    return False


def parse_template(docx_path, project_name=None):
    """Parse a Word template and extract its structure."""
    doc = Document(docx_path)
    sections = []
    current_section = None
    doc_title = None

    # Extract document title from first paragraph if it looks like a title
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = (para.style.name or '').lower()
            if 'title' in style_name or 'heading' in style_name:
                doc_title = text
            break

    # Build section list from headings
    para_list = list(doc.paragraphs)
    i = 0
    while i < len(para_list):
        para = para_list[i]
        text = para.text.strip()
        style_name = para.style.name or ''

        if is_heading(para) and text:
            # Save previous section
            if current_section:
                sections.append(current_section)

            level = get_heading_level(para)
            current_section = {
                "heading": text,
                "level": level,
                "style": style_name,
                "content": [],
                "tables": 0,
                "subsections": []
            }
            if doc_title is None and level == 1:
                doc_title = text
        elif current_section and text:
            current_section["content"].append({
                "index": i,
                "text": text,
                "style": style_name,
                "is_placeholder": is_placeholder(text)
            })
        elif current_section and not text:
            # Empty paragraph — could be a content slot to fill
            current_section["content"].append({
                "index": i,
                "text": "",
                "style": style_name,
                "is_placeholder": True
            })
        i += 1

    # Save last section
    if current_section:
        sections.append(current_section)

    # Count tables and images
    table_count = len(doc.tables)
    image_count = sum(1 for para in doc.paragraphs if has_images(para))

    # Build hierarchical structure
    def build_hierarchy(flat_sections, parent_level=0):
        result = []
        i = 0
        while i < len(flat_sections):
            sec = flat_sections[i]
            if sec["level"] > parent_level:
                # Find all subsections
                children = []
                j = i + 1
                while j < len(flat_sections):
                    if flat_sections[j]["level"] <= sec["level"]:
                        break
                    children.append(flat_sections[j])
                    j += 1
                if children:
                    sec["subsections"] = build_hierarchy(children, sec["level"])
                result.append(sec)
                i = j
            else:
                i += 1
        return result

    # Group by top-level headings
    top_level = min((s["level"] for s in sections), default=1)
    hierarchy = build_hierarchy(sections, top_level - 1)

    output = {
        "source": str(docx_path),
        "title": doc_title,
        "sections": hierarchy,
        "flat_sections": sections,
        "metadata": {
            "total_headings": len(sections),
            "total_paragraphs": len(para_list),
            "total_tables": table_count,
            "has_images": image_count > 0,
            "image_count": image_count
        }
    }

    # Add suggested title if project name is provided
    if project_name:
        suggested_title = _generate_title_from_project_name(doc, project_name)
        output["suggested_title"] = suggested_title

    return output


def convert_doc_to_docx(doc_path):
    """Convert .doc to .docx using textutil (macOS) or libreoffice (Linux).

    Returns: (converted_path, format_lost)
      - converted_path: path to .docx file
      - format_lost: True if formatting/styles were lost in conversion
    """
    import subprocess, tempfile
    doc_path = str(doc_path)
    if doc_path.lower().endswith('.docx'):
        return doc_path, False
    if not doc_path.lower().endswith('.doc'):
        return doc_path, False

    # Try macOS textutil
    try:
        tmp_dir = tempfile.mkdtemp()
        result = subprocess.run(
            ['textutil', '-convert', 'docx', '-output',
             os.path.join(tmp_dir, 'converted.docx'), doc_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            converted = os.path.join(tmp_dir, 'converted.docx')
            # Check if headings were preserved
            try:
                from docx import Document
                test_doc = Document(converted)
                has_headings = any(
                    'heading' in (p.style.name or '').lower()
                    for p in test_doc.paragraphs
                )
                if not has_headings:
                    print(f"WARNING: .doc converted but heading styles lost. \
Agent should read template text and generate new .docx.", file=sys.stderr)
                    return converted, True
            except Exception:
                pass
            print(f"Converted .doc → .docx via textutil", file=sys.stderr)
            return converted, False
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Try libreoffice
    try:
        tmp_dir = tempfile.mkdtemp()
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx',
             '--outdir', tmp_dir, doc_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            base = Path(doc_path).stem + '.docx'
            converted = os.path.join(tmp_dir, base)
            if os.path.exists(converted):
                print(f"Converted .doc → .docx via libreoffice", file=sys.stderr)
                return converted, False
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    print(f"WARNING: Cannot convert .doc file.", file=sys.stderr)
    return doc_path, True


def parse_doc_text_structure(docx_path):
    """Parse a converted .doc file where heading styles were lost.

    Extracts text content and attempts to identify sections by patterns
    (e.g., numbered headings like '1 变更目的', or known template keywords).
    """
    from docx import Document

    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Common section keywords for change documents
    CHANGE_SECTIONS = [
        '变更方案', '变更计划', '变更目的', '变更内容', '变更前准备',
        '变更执行', '变更步骤', '升级完成后检查', '变更回退',
        '参数设置', '配置更新', '软件及文档更新', '网络配置',
        '硬件配置', '变更后工作', '变更后检查', '部署架构图',
        '版本', '编制', '校对', '部门', '日期'
    ]

    # Try to identify sections by content
    sections = []
    current = None
    for text in paragraphs:
        # Check if this looks like a section header
        is_section = False
        for kw in CHANGE_SECTIONS:
            if kw in text and len(text) < 30:
                is_section = True
                break

        # Also check for numbered headings: 1 xxx, 2.1 xxx, etc.
        if re.match(r'^\d+[\.\s]', text) and len(text) < 50:
            is_section = True

        if is_section:
            if current:
                sections.append(current)
            current = {
                "heading": text,
                "level": 1,
                "style": "Normal",
                "content": [],
                "tables": 0,
                "subsections": []
            }
        elif current:
            current["content"].append({
                "index": len(sections),
                "text": text,
                "style": "Normal",
                "is_placeholder": False
            })

    if current:
        sections.append(current)

    return {
        "source": str(docx_path),
        "title": paragraphs[0] if paragraphs else None,
        "sections": sections,
        "flat_sections": sections,
        "all_paragraphs": paragraphs,
        "metadata": {
            "total_headings": len(sections),
            "total_paragraphs": len(paragraphs),
            "total_tables": len(doc.tables),
            "has_images": False,
            "image_count": 0
        }
    }


def main():
    parser = argparse.ArgumentParser(description="Parse Word template structure")
    parser.add_argument("template", help="Path to the .doc or .docx template file")
    parser.add_argument("--output", "-o", help="Output JSON file (default: stdout)")
    parser.add_argument("--flat", action="store_true", help="Output flat section list only")
    parser.add_argument("--project-name", help="Project name for smart title generation")

    args = parser.parse_args()
    template_path = Path(args.template)

    if not template_path.exists():
        print(f"ERROR: File not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    # Auto-convert .doc to .docx
    resolved_path, format_lost = convert_doc_to_docx(template_path)
    template_path = Path(resolved_path)

    if format_lost:
        # For .doc files with lost formatting, extract text structure instead
        result = parse_doc_text_structure(template_path)
        result["format_lost"] = True
        result["note"] = "Original .doc formatting lost. Agent should generate new .docx from this structure."
    else:
        result = parse_template(template_path, project_name=args.project_name)

    if args.flat:
        result = result["flat_sections"]

    json_str = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        Path(args.output).write_text(json_str, encoding='utf-8')
        print(f"Structure saved to {args.output}", file=sys.stderr)
    else:
        print(json_str)


if __name__ == "__main__":
    main()
